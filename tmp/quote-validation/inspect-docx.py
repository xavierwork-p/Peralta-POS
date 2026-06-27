from pathlib import Path
from zipfile import ZipFile
import json
import xml.etree.ElementTree as ET

from docx import Document

path = Path("tmp/quote-validation/docs/quote-demo.docx")
doc = Document(path)
texts = [p.text for p in doc.paragraphs if p.text.strip()]
tables = []

for table in doc.tables:
    tables.append({
        "rows": len(table.rows),
        "cols": len(table.columns),
        "first_row": [cell.text for cell in table.rows[0].cells] if table.rows else [],
    })

with ZipFile(path) as archive:
    xml = archive.read("word/document.xml")

root = ET.fromstring(xml)
ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
section = root.find(".//w:sectPr", ns)
page_size = section.find("w:pgSz", ns) if section is not None else None
margins = section.find("w:pgMar", ns) if section is not None else None

print(json.dumps({
    "paragraphs": texts[:12],
    "table_count": len(tables),
    "tables": tables,
    "has_section_properties": section is not None,
    "page_size": page_size.attrib if page_size is not None else None,
    "margins": margins.attrib if margins is not None else None,
}, ensure_ascii=False, indent=2))
